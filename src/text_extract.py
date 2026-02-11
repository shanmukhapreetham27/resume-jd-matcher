from __future__ import annotations

from io import BytesIO
import re

from docx import Document
from pypdf import PdfReader


def clean_text(text: str) -> str:
    text = text.replace("\x00", " ")
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


def extract_text_from_pdf_bytes(file_bytes: bytes) -> str:
    reader = PdfReader(BytesIO(file_bytes))
    pages = [page.extract_text() or "" for page in reader.pages]
    return clean_text("\n".join(pages))


def extract_text_from_docx_bytes(file_bytes: bytes) -> str:
    doc = Document(BytesIO(file_bytes))
    parts = [p.text for p in doc.paragraphs if p.text.strip()]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    parts.append(cell_text)

    return clean_text("\n".join(parts))


def extract_text_from_txt_bytes(file_bytes: bytes) -> str:
    return clean_text(file_bytes.decode("utf-8", errors="ignore"))


def extract_resume_text(uploaded_file) -> str:
    if uploaded_file is None:
        raise ValueError("No file uploaded.")

    suffix = uploaded_file.name.lower().split(".")[-1]
    file_bytes = uploaded_file.getvalue()
    if not file_bytes:
        raise ValueError("Uploaded file is empty.")

    if suffix == "pdf":
        return extract_text_from_pdf_bytes(file_bytes)
    if suffix == "docx":
        return extract_text_from_docx_bytes(file_bytes)
    if suffix == "txt":
        return extract_text_from_txt_bytes(file_bytes)

    raise ValueError("Unsupported file type. Use PDF, DOCX, or TXT.")
